import base64
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from vaf.core.docx_export import export_native_docx
from vaf.core.docx_import import import_docx_to_native_model
from vaf.core.docx_native_model import NativeDocxDocument, flatten_document_text


SAMPLE_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9l9Y4AAAAASUVORK5CYII="
)


def test_native_docx_document_roundtrip_dict_and_flatten():
    payload = {
        "schema_version": 1,
        "source_format": "docx",
        "title": "Contract",
        "path": "/tmp/contract.docx",
        "warnings": [{"id": "w1", "code": "info", "message": "ok", "severity": "info"}],
        "sections": [
            {
                "id": "section-0",
                "properties": {"start_type": "newPage"},
                "header": {"paragraphs": []},
                "footer": {"paragraphs": []},
                "blocks": [
                    {
                        "id": "paragraph-0",
                        "type": "paragraph",
                        "style_name": "Heading 1",
                        "alignment": "left",
                        "list_kind": "none",
                        "list_level": 0,
                        "page_break_before": False,
                        "keep_with_next": False,
                        "keep_together": False,
                        "runs": [{"id": "run-0", "text": "Clause A", "bold": True}],
                    },
                    {
                        "id": "table-0",
                        "type": "table",
                        "style_name": "Table Grid",
                        "rows": [
                            {
                                "id": "row-0",
                                "cells": [
                                    {
                                        "id": "cell-0",
                                        "column_span": 1,
                                        "row_span": 1,
                                        "paragraphs": [
                                            {
                                                "id": "cell-p-0",
                                                "type": "paragraph",
                                                "style_name": "Normal",
                                                "alignment": "left",
                                                "list_kind": "none",
                                                "list_level": 0,
                                                "page_break_before": False,
                                                "keep_with_next": False,
                                                "keep_together": False,
                                                "runs": [{"id": "cell-run-0", "text": "Cell text"}],
                                            }
                                        ],
                                    }
                                ],
                            }
                        ],
                    },
                ],
            }
        ],
    }

    document = NativeDocxDocument.from_dict(payload)

    assert document.title == "Contract"
    assert document.sections[0].blocks[0].type == "paragraph"
    assert document.sections[0].blocks[1].type == "table"
    assert flatten_document_text(document) == "Clause A\nCell text"
    assert document.to_dict()["title"] == "Contract"


def test_import_docx_to_native_model_preserves_block_order_and_structure(tmp_path: Path):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.core_properties.title = "Native Import"
    doc.sections[0].header.paragraphs[0].text = "Header text"
    doc.sections[0].footer.paragraphs[0].text = "Footer text"
    doc.add_heading("Overview", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Styled text")
    run.bold = True
    doc.add_paragraph("First bullet", style="List Bullet")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Left"
    table.rows[0].cells[1].text = "Right"
    doc.save(source)

    native = import_docx_to_native_model(source)

    assert native.path == str(source.resolve())
    assert native.sections[0].header.paragraphs[0].runs[0].text == "Header text"
    assert native.sections[0].footer.paragraphs[0].runs[0].text == "Footer text"
    assert native.sections[0].blocks[0].type == "paragraph"
    assert native.sections[0].blocks[0].style_name == "Heading 1"
    assert native.sections[0].blocks[1].type == "paragraph"
    assert native.sections[0].blocks[1].runs[0].bold is True
    assert native.sections[0].blocks[2].type == "paragraph"
    assert native.sections[0].blocks[2].list_kind == "bullet"
    assert native.sections[0].blocks[3].type == "table"


def test_import_docx_to_native_model_treats_ooxml_zero_on_off_flags_as_false(tmp_path: Path):
    source = tmp_path / "flags-zero.docx"
    doc = Document()
    paragraph = doc.add_paragraph("Paragraph with disabled pagination flags")
    p_pr = paragraph._element.get_or_add_pPr()
    for tag in ("w:pageBreakBefore", "w:keepNext", "w:keepLines"):
        element = OxmlElement(tag)
        element.set(qn("w:val"), "0")
        p_pr.append(element)
    doc.save(source)

    native = import_docx_to_native_model(source)
    block = native.sections[0].blocks[0]

    assert block.type == "paragraph"
    assert block.page_break_before is False
    assert block.keep_with_next is False
    assert block.keep_together is False


def test_export_native_docx_writes_sections_lists_tables_and_images(tmp_path: Path):
    target = tmp_path / "exported.docx"
    document = NativeDocxDocument.from_dict(
        {
            "schema_version": 1,
            "source_format": "docx",
            "title": "Native Export",
            "path": str(target),
            "warnings": [],
            "sections": [
                {
                    "id": "section-0",
                    "properties": {
                        "start_type": "newPage",
                        "margin_top_twips": 1440,
                        "margin_right_twips": 1440,
                        "margin_bottom_twips": 1440,
                        "margin_left_twips": 1440,
                    },
                    "header": {
                        "paragraphs": [
                            {
                                "id": "header-0",
                                "type": "paragraph",
                                "style_name": "Normal",
                                "alignment": "left",
                                "list_kind": "none",
                                "list_level": 0,
                                "page_break_before": False,
                                "keep_with_next": False,
                                "keep_together": False,
                                "runs": [{"id": "header-run-0", "text": "Header line"}],
                            }
                        ]
                    },
                    "footer": {
                        "paragraphs": [
                            {
                                "id": "footer-0",
                                "type": "paragraph",
                                "style_name": "Normal",
                                "alignment": "left",
                                "list_kind": "none",
                                "list_level": 0,
                                "page_break_before": False,
                                "keep_with_next": False,
                                "keep_together": False,
                                "runs": [{"id": "footer-run-0", "text": "Footer line"}],
                            }
                        ]
                    },
                    "blocks": [
                        {
                            "id": "heading-0",
                            "type": "paragraph",
                            "style_name": "Heading 1",
                            "alignment": "left",
                            "list_kind": "none",
                            "list_level": 0,
                            "page_break_before": False,
                            "keep_with_next": False,
                            "keep_together": False,
                            "runs": [{"id": "heading-run-0", "text": "Overview"}],
                        },
                        {
                            "id": "bullet-0",
                            "type": "paragraph",
                            "style_name": "Normal",
                            "alignment": "left",
                            "list_kind": "bullet",
                            "list_level": 0,
                            "page_break_before": False,
                            "keep_with_next": False,
                            "keep_together": False,
                            "runs": [{"id": "bullet-run-0", "text": "Bullet item"}],
                        },
                        {
                            "id": "table-0",
                            "type": "table",
                            "style_name": "Table Grid",
                            "rows": [
                                {
                                    "id": "row-0",
                                    "cells": [
                                        {
                                            "id": "cell-0",
                                            "column_span": 1,
                                            "row_span": 1,
                                            "paragraphs": [
                                                {
                                                    "id": "cell-0-p-0",
                                                    "type": "paragraph",
                                                    "style_name": "Normal",
                                                    "alignment": "left",
                                                    "list_kind": "none",
                                                    "list_level": 0,
                                                    "page_break_before": False,
                                                    "keep_with_next": False,
                                                    "keep_together": False,
                                                    "runs": [{"id": "cell-run-0", "text": "Left"}],
                                                }
                                            ],
                                        },
                                        {
                                            "id": "cell-1",
                                            "column_span": 1,
                                            "row_span": 1,
                                            "paragraphs": [
                                                {
                                                    "id": "cell-1-p-0",
                                                    "type": "paragraph",
                                                    "style_name": "Normal",
                                                    "alignment": "left",
                                                    "list_kind": "none",
                                                    "list_level": 0,
                                                    "page_break_before": False,
                                                    "keep_with_next": False,
                                                    "keep_together": False,
                                                    "runs": [{"id": "cell-run-1", "text": "Right"}],
                                                }
                                            ],
                                        },
                                    ],
                                }
                            ],
                        },
                        {
                            "id": "image-0",
                            "type": "image",
                            "alt_text": "Logo",
                            "filename": "logo.png",
                            "content_type": "image/png",
                            "base64_data": SAMPLE_PNG_BASE64,
                            "width_px": 16,
                            "height_px": 16,
                            "anchor_kind": "inline",
                        },
                    ],
                }
            ],
        }
    )

    export_native_docx(document, target)

    saved = Document(str(target))
    texts = [paragraph.text for paragraph in saved.paragraphs if paragraph.text]

    assert saved.core_properties.title == "Native Export"
    assert "Overview" in texts
    assert "Bullet item" in texts
    assert "Logo" in texts
    assert saved.sections[0].header.paragraphs[0].text == "Header line"
    assert saved.sections[0].footer.paragraphs[0].text == "Footer line"
    assert saved.tables[0].rows[0].cells[0].text == "Left"
    assert saved.tables[0].rows[0].cells[1].text == "Right"


def test_import_export_roundtrip_keeps_supported_text_features(tmp_path: Path):
    source = tmp_path / "roundtrip_source.docx"
    exported = tmp_path / "roundtrip_exported.docx"

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Header"
    doc.add_heading("Title", level=1)
    p = doc.add_paragraph()
    p.add_run("Hello ").bold = True
    p.add_run("world")
    doc.add_paragraph("Numbered entry", style="List Number")
    doc.save(source)

    native = import_docx_to_native_model(source)
    export_native_docx(native, exported)
    reopened = import_docx_to_native_model(exported)

    assert flatten_document_text(reopened).startswith("Header\nTitle\nHello world")
    assert any(
        block.type == "paragraph" and getattr(block, "list_kind", "none") == "numbered"
        for block in reopened.sections[0].blocks
    )


def test_import_drops_trailing_body_sectpr_without_warning(tmp_path):
    """The trailing body <w:sectPr> (the document's final section properties) must NOT be imported as a
    read-only 'Unsupported' placeholder — it is captured via doc.sections and re-emitted on export, so it
    should produce neither a placeholder block nor the spurious warning."""
    source = tmp_path / "trailing-sectpr.docx"
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Body text")
    doc.save(source)

    native = import_docx_to_native_model(source)

    assert all(block.type != "unsupported" for section in native.sections for block in section.blocks)
    assert not any(w.code == "unsupported_block" for w in native.warnings)


def test_trailing_sectpr_roundtrips_section_properties(tmp_path):
    """Dropping the trailing sectPr placeholder loses no section data: import -> export -> re-import keeps
    the section properties and still produces no unsupported block / warning."""
    source = tmp_path / "rt-src.docx"
    exported = tmp_path / "rt-out.docx"
    doc = Document()
    doc.add_paragraph("Body")
    doc.save(source)

    native = import_docx_to_native_model(source)
    export_native_docx(native, exported)
    reopened = import_docx_to_native_model(exported)

    assert reopened.sections[0].properties.page_width_twips is not None
    assert all(block.type != "unsupported" for section in reopened.sections for block in section.blocks)
    assert not any(w.code == "unsupported_block" for w in reopened.warnings)
