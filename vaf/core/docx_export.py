from __future__ import annotations

import base64
import os
import shutil
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches, Pt

from vaf.core.docx_native_model import (
    DocxHeaderFooter,
    DocxImage,
    DocxPageBreak,
    DocxParagraph,
    DocxSection,
    DocxTable,
    DocxUnsupportedBlock,
    NativeDocxDocument,
)


_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_SECTION_START_MAP = {
    "newPage": WD_SECTION_START.NEW_PAGE,
    "continuous": WD_SECTION_START.CONTINUOUS,
    "evenPage": WD_SECTION_START.EVEN_PAGE,
    "oddPage": WD_SECTION_START.ODD_PAGE,
    "newColumn": WD_SECTION_START.NEW_COLUMN,
}


def export_native_docx(document_model: NativeDocxDocument, file_path: str | Path) -> Path:
    target = Path(file_path).resolve()
    target.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _apply_title(doc, document_model.title)

    if document_model.sections:
        _apply_section_properties(doc.sections[0], document_model.sections[0])
        _write_header_footer(doc.sections[0].header, document_model.sections[0].header)
        _write_header_footer(doc.sections[0].footer, document_model.sections[0].footer)

    for section_index, section in enumerate(document_model.sections):
        if section_index > 0:
            new_section = doc.add_section(_SECTION_START_MAP.get(section.properties.start_type, WD_SECTION_START.NEW_PAGE))
            _apply_section_properties(new_section, section)
            _write_header_footer(new_section.header, section.header)
            _write_header_footer(new_section.footer, section.footer)
        for block in section.blocks:
            _write_block(doc, block)

    fd, tmp_path = tempfile.mkstemp(prefix="vaf_docx_native_", suffix=".docx", dir=str(target.parent))
    try:
        os.close(fd)
        doc.save(tmp_path)
        if target.exists():
            target.unlink()
        shutil.move(tmp_path, str(target))
    finally:
        if os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    return target


def _apply_title(doc: Document, title: str) -> None:
    if title.strip():
        doc.core_properties.title = title


def _apply_section_properties(section_ref, section: DocxSection) -> None:
    props = section.properties
    if props.page_width_twips is not None:
        section_ref.page_width = _twips_to_emu(props.page_width_twips)
    if props.page_height_twips is not None:
        section_ref.page_height = _twips_to_emu(props.page_height_twips)
    if props.margin_top_twips is not None:
        section_ref.top_margin = _twips_to_emu(props.margin_top_twips)
    if props.margin_right_twips is not None:
        section_ref.right_margin = _twips_to_emu(props.margin_right_twips)
    if props.margin_bottom_twips is not None:
        section_ref.bottom_margin = _twips_to_emu(props.margin_bottom_twips)
    if props.margin_left_twips is not None:
        section_ref.left_margin = _twips_to_emu(props.margin_left_twips)


def _twips_to_emu(value: int):
    return Inches(value / 1440.0)


def _write_header_footer(container, header_footer: DocxHeaderFooter) -> None:
    paragraphs = header_footer.paragraphs
    if not paragraphs:
        return
    if container.paragraphs:
        first = container.paragraphs[0]
        if not first.text and len(first.runs) == 0:
            container._element.remove(first._element)
    for paragraph in paragraphs:
        _write_paragraph(container, paragraph)


def _write_block(doc: Document, block) -> None:
    if isinstance(block, DocxParagraph):
        _write_paragraph(doc, block)
        return
    if isinstance(block, DocxTable):
        _write_table(doc, block)
        return
    if isinstance(block, DocxImage):
        _write_image(doc, block)
        return
    if isinstance(block, DocxPageBreak):
        doc.add_page_break()
        return
    if isinstance(block, DocxUnsupportedBlock):
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"[Unsupported content: {block.label}]")


def _write_paragraph(container, paragraph_model: DocxParagraph):
    style_name = paragraph_model.style_name or "Normal"
    list_kind = paragraph_model.list_kind or "none"
    if list_kind == "bullet":
        style_name = "List Bullet"
    elif list_kind == "numbered":
        style_name = "List Number"

    paragraph = container.add_paragraph(style=style_name if style_name else None)
    paragraph.alignment = _ALIGNMENT_MAP.get(paragraph_model.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.paragraph_format.keep_with_next = paragraph_model.keep_with_next
    paragraph.paragraph_format.keep_together = paragraph_model.keep_together
    paragraph.paragraph_format.page_break_before = paragraph_model.page_break_before

    for run_model in paragraph_model.runs:
        run = paragraph.add_run(run_model.text)
        run.bold = run_model.bold
        run.italic = run_model.italic
        run.underline = run_model.underline
        if run_model.font_name:
            run.font.name = run_model.font_name
        if run_model.font_size_pt is not None:
            run.font.size = Pt(run_model.font_size_pt)
        if run_model.color:
            cleaned = run_model.color.replace("#", "")
            if len(cleaned) == 6:
                run.font.color.rgb = RGBColor.from_string(cleaned)


def _write_table(doc: Document, table_model: DocxTable) -> None:
    row_count = max(len(table_model.rows), 1)
    col_count = max((len(row.cells) for row in table_model.rows), default=1)
    table = doc.add_table(rows=row_count, cols=col_count)
    try:
        table.style = table_model.style_name or "Table Grid"
    except Exception:
        table.style = "Table Grid"

    for row_index, row_model in enumerate(table_model.rows):
        for cell_index, cell_model in enumerate(row_model.cells):
            if cell_index >= len(table.rows[row_index].cells):
                continue
            cell = table.rows[row_index].cells[cell_index]
            _clear_cell(cell)
            for paragraph_model in cell_model.paragraphs:
                _write_paragraph(cell, paragraph_model)


def _clear_cell(cell) -> None:
    for paragraph in list(cell.paragraphs):
        if paragraph._element is not None:
            cell._element.remove(paragraph._element)


def _write_image(doc: Document, image_model: DocxImage) -> None:
    paragraph = doc.add_paragraph()
    if image_model.base64_data:
        data = BytesIO(base64.b64decode(image_model.base64_data))
        width_inches = image_model.width_px / 96.0 if image_model.width_px else None
        if width_inches:
            paragraph.add_run().add_picture(data, width=Inches(width_inches))
        else:
            paragraph.add_run().add_picture(data)
    if image_model.alt_text:
        paragraph.add_run(image_model.alt_text)
