from __future__ import annotations

import os
import re
import shutil
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable


@dataclass
class DocumentBlock:
    """A normalized semantic block inside a document section."""

    type: str
    text: str = ""
    items: list[str] = field(default_factory=list)


@dataclass
class DocumentSection:
    """A normalized document section."""

    title: str
    heading_level: int = 2
    blocks: list[DocumentBlock] = field(default_factory=list)


@dataclass
class DocumentModel:
    """Canonical intermediate representation for document rendering."""

    title: str
    document_type: str
    sections: list[DocumentSection] = field(default_factory=list)


_HEADING_MARKER_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_UNDERLINE_RE = re.compile(r"^(?:={3,}|-{3,})$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^(?:\d+|[a-zA-Z])[\.\)]\s+(.*)$")


def sanitize_inline_text(text: str) -> str:
    """Remove decorative markdown and normalize inline whitespace."""

    if not text:
        return ""

    cleaned = str(text).replace("\r\n", "\n").replace("\r", "\n")
    cleaned = re.sub(r"```[\s\S]*?```", "", cleaned)
    cleaned = re.sub(r"(?m)^\s*#{1,6}\s+", "", cleaned)
    cleaned = re.sub(r"(?m)^\s*(?:={3,}|-{3,})\s*$", "", cleaned)
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__(.+?)__", r"\1", cleaned)
    cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in cleaned.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines).strip()


def _normalize_paragraph_text(lines: Iterable[str]) -> str:
    return sanitize_inline_text(" ".join(line.strip() for line in lines if line.strip()))


def _strip_list_marker(line: str) -> str:
    line = line.strip()
    bullet_match = _BULLET_RE.match(line)
    if bullet_match:
        return sanitize_inline_text(bullet_match.group(1))
    numbered_match = _NUMBERED_RE.match(line)
    if numbered_match:
        return sanitize_inline_text(numbered_match.group(1))
    return sanitize_inline_text(line)


def _split_into_chunks(text: str) -> list[list[str]]:
    chunks: list[list[str]] = []
    current: list[str] = []

    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        if line.strip():
            current.append(line.rstrip())
        elif current:
            chunks.append(current)
            current = []

    if current:
        chunks.append(current)

    return chunks


def blocks_from_text(text: str) -> list[DocumentBlock]:
    """Parse plain text or light markdown into normalized blocks."""

    blocks: list[DocumentBlock] = []
    for chunk in _split_into_chunks(text):
        stripped = [line.strip() for line in chunk if line.strip()]
        if not stripped:
            continue

        if all(_BULLET_RE.match(line) for line in stripped):
            items = [_strip_list_marker(line) for line in stripped]
            items = [item for item in items if item]
            if items:
                blocks.append(DocumentBlock(type="bullet_list", items=items))
                continue

        if all(_NUMBERED_RE.match(line) for line in stripped):
            items = [_strip_list_marker(line) for line in stripped]
            items = [item for item in items if item]
            if items:
                blocks.append(DocumentBlock(type="numbered_list", items=items))
                continue

        paragraph = _normalize_paragraph_text(stripped)
        if paragraph:
            blocks.append(DocumentBlock(type="paragraph", text=paragraph))

    return blocks


def coerce_block(payload: Any) -> DocumentBlock | None:
    """Convert raw model output or legacy text into a normalized block."""

    if isinstance(payload, DocumentBlock):
        if payload.type == "paragraph":
            text = sanitize_inline_text(payload.text)
            return DocumentBlock(type="paragraph", text=text) if text else None
        items = [sanitize_inline_text(item) for item in payload.items]
        items = [item for item in items if item]
        return DocumentBlock(type=payload.type, items=items) if items else None

    if isinstance(payload, str):
        text = sanitize_inline_text(payload)
        return DocumentBlock(type="paragraph", text=text) if text else None

    if not isinstance(payload, dict):
        return None

    block_type = str(payload.get("type", "paragraph")).strip().lower()
    if block_type not in {"paragraph", "bullet_list", "numbered_list"}:
        block_type = "paragraph"

    if block_type == "paragraph":
        text = sanitize_inline_text(str(payload.get("text", "")).strip())
        if text:
            return DocumentBlock(type="paragraph", text=text)
        return None

    raw_items = payload.get("items", [])
    if not isinstance(raw_items, list):
        raw_items = [raw_items]

    items = [sanitize_inline_text(str(item)) for item in raw_items]
    items = [item for item in items if item]
    if items:
        return DocumentBlock(type=block_type, items=items)

    fallback_text = sanitize_inline_text(str(payload.get("text", "")))
    if fallback_text:
        parsed = blocks_from_text(fallback_text)
        for block in parsed:
            if block.type == block_type:
                return block
    return None


def coerce_section(payload: Any, fallback_title: str = "Section", default_level: int = 2) -> DocumentSection:
    """Convert raw section payloads into a normalized section model."""

    if isinstance(payload, DocumentSection):
        normalized_blocks = [coerce_block(block) for block in payload.blocks]
        blocks = [block for block in normalized_blocks if block]
        if not blocks:
            blocks = [DocumentBlock(type="paragraph", text="Content unavailable.")]
        level = max(2, min(6, int(payload.heading_level or default_level)))
        title = sanitize_inline_text(payload.title) or fallback_title
        return DocumentSection(title=title, heading_level=level, blocks=blocks)

    if isinstance(payload, str):
        blocks = blocks_from_text(payload)
        if not blocks:
            blocks = [DocumentBlock(type="paragraph", text="Content unavailable.")]
        return DocumentSection(
            title=sanitize_inline_text(fallback_title) or "Section",
            heading_level=max(2, min(6, default_level)),
            blocks=blocks,
        )

    if not isinstance(payload, dict):
        return DocumentSection(
            title=sanitize_inline_text(fallback_title) or "Section",
            heading_level=max(2, min(6, default_level)),
            blocks=[DocumentBlock(type="paragraph", text="Content unavailable.")],
        )

    title = sanitize_inline_text(str(payload.get("title", fallback_title))) or fallback_title

    try:
        heading_level = int(payload.get("heading_level", default_level))
    except (TypeError, ValueError):
        heading_level = default_level
    heading_level = max(2, min(6, heading_level))

    normalized_blocks: list[DocumentBlock] = []
    raw_blocks = payload.get("blocks", [])
    if isinstance(raw_blocks, list):
        for raw_block in raw_blocks:
            block = coerce_block(raw_block)
            if block:
                normalized_blocks.append(block)

    if not normalized_blocks:
        fallback_content = payload.get("content") or payload.get("text") or payload.get("body") or ""
        if isinstance(fallback_content, list):
            fallback_content = "\n".join(str(item) for item in fallback_content)
        normalized_blocks = blocks_from_text(str(fallback_content))

    if not normalized_blocks:
        normalized_blocks = [DocumentBlock(type="paragraph", text="Content unavailable.")]

    return DocumentSection(title=title, heading_level=heading_level, blocks=normalized_blocks)


def build_document_model(title: str, document_type: str, sections: list[Any]) -> DocumentModel:
    """Build a normalized document model from raw sections."""

    normalized_title = sanitize_inline_text(title) or "Document"
    normalized_type = sanitize_inline_text(document_type) or "document"
    normalized_sections: list[DocumentSection] = []

    for index, section in enumerate(sections, start=1):
        normalized_sections.append(
            coerce_section(section, fallback_title=f"Section {index}", default_level=2)
        )

    if not normalized_sections:
        normalized_sections = [
            DocumentSection(
                title="Content",
                heading_level=2,
                blocks=[DocumentBlock(type="paragraph", text="Content unavailable.")],
            )
        ]

    return DocumentModel(
        title=normalized_title,
        document_type=normalized_type,
        sections=normalized_sections,
    )


def _convert_underlined_headings(text: str) -> tuple[str, str | None]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    converted: list[str] = []
    derived_title: str | None = None
    index = 0

    while index < len(lines):
        current = lines[index].rstrip()
        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""
        if current.strip() and _UNDERLINE_RE.fullmatch(next_line):
            if derived_title is None and not converted:
                derived_title = sanitize_inline_text(current)
            else:
                converted.append(f"## {current.strip()}")
            index += 2
            continue

        converted.append(current)
        index += 1

    return "\n".join(converted), derived_title


def infer_document_model(title: str, document_type: str, content: str) -> DocumentModel:
    """Infer the canonical document model from markdown or plain text content."""

    transformed, derived_title = _convert_underlined_headings(content)
    lines = transformed.split("\n")
    resolved_title = sanitize_inline_text(title) or derived_title or "Document"

    sections: list[DocumentSection] = []
    preface_lines: list[str] = []
    current_title: str | None = None
    current_level = 2
    current_lines: list[str] = []

    def flush_current() -> None:
        nonlocal current_title, current_level, current_lines
        if current_title is None:
            return
        blocks = blocks_from_text("\n".join(current_lines))
        if not blocks:
            blocks = [DocumentBlock(type="paragraph", text="Content unavailable.")]
        sections.append(
            DocumentSection(title=current_title, heading_level=current_level, blocks=blocks)
        )
        current_title = None
        current_level = 2
        current_lines = []

    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped and current_title is None:
            preface_lines.append(raw_line)
            continue

        heading_match = _HEADING_MARKER_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = sanitize_inline_text(heading_match.group(2))
            if level == 1 and heading_text and not sections and current_title is None and not any(
                line.strip() for line in preface_lines
            ):
                resolved_title = heading_text
                continue

            if level >= 2 and heading_text:
                flush_current()
                current_title = heading_text
                current_level = max(2, min(6, level))
                continue

        if current_title is None:
            preface_lines.append(raw_line)
        else:
            current_lines.append(raw_line)

    flush_current()

    preface_text = "\n".join(preface_lines).strip()
    if preface_text:
        preface_blocks = blocks_from_text(preface_text)
        if preface_blocks:
            sections.insert(
                0,
                DocumentSection(title="Content", heading_level=2, blocks=preface_blocks),
            )

    if not sections:
        sections = [
            DocumentSection(
                title="Content",
                heading_level=2,
                blocks=blocks_from_text(content) or [DocumentBlock(type="paragraph", text="Content unavailable.")],
            )
        ]

    return DocumentModel(title=resolved_title, document_type=document_type, sections=sections)


def render_markdown(model: DocumentModel) -> str:
    """Render the canonical model to stable markdown."""

    parts: list[str] = [f"# {sanitize_inline_text(model.title)}"]

    for section in model.sections:
        heading_level = max(2, min(6, section.heading_level))
        parts.append("")
        parts.append(f'{"#" * heading_level} {sanitize_inline_text(section.title)}')
        parts.append("")

        for block in section.blocks:
            if block.type == "paragraph":
                parts.append(sanitize_inline_text(block.text))
            elif block.type == "bullet_list":
                parts.extend(f"- {sanitize_inline_text(item)}" for item in block.items)
            elif block.type == "numbered_list":
                parts.extend(f"{index}. {sanitize_inline_text(item)}" for index, item in enumerate(block.items, start=1))
            parts.append("")

    return "\n".join(parts).strip() + "\n"


def _html_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def render_section_html(section: DocumentSection) -> str:
    """Render ONE section to an HTML fragment for the live document viewer.

    Mirrors render_markdown's block handling (paragraph / bullet_list /
    numbered_list). Placeholders like ``{{NAME}}`` are kept verbatim (the frontend
    highlights them); text is HTML-escaped, structure is minimal."""
    level = max(2, min(6, section.heading_level))
    parts: list[str] = [f"<h{level}>{_html_escape(sanitize_inline_text(section.title))}</h{level}>"]
    for block in section.blocks:
        if block.type == "paragraph":
            txt = _html_escape(sanitize_inline_text(block.text))
            if txt:
                parts.append(f"<p>{txt}</p>")
        elif block.type == "bullet_list":
            items = "".join(f"<li>{_html_escape(sanitize_inline_text(it))}</li>" for it in block.items if it)
            if items:
                parts.append(f"<ul>{items}</ul>")
        elif block.type == "numbered_list":
            items = "".join(f"<li>{_html_escape(sanitize_inline_text(it))}</li>" for it in block.items if it)
            if items:
                parts.append(f"<ol>{items}</ol>")
    return "".join(parts)


def render_text(model: DocumentModel) -> str:
    """Render the canonical model to plain text."""

    parts: list[str] = [sanitize_inline_text(model.title), "=" * len(sanitize_inline_text(model.title)), ""]

    for section in model.sections:
        title = sanitize_inline_text(section.title)
        parts.extend([title, "-" * len(title), ""])
        for block in section.blocks:
            if block.type == "paragraph":
                parts.extend([sanitize_inline_text(block.text), ""])
            elif block.type == "bullet_list":
                parts.extend([f"- {sanitize_inline_text(item)}" for item in block.items])
                parts.append("")
            elif block.type == "numbered_list":
                parts.extend(
                    f"{index}. {sanitize_inline_text(item)}"
                    for index, item in enumerate(block.items, start=1)
                )
                parts.append("")

    return "\n".join(parts).strip() + "\n"


def estimate_document_length(model: DocumentModel) -> int:
    """Use normalized markdown length as a stable size estimate."""

    return len(render_markdown(model))


def save_document_model_as_docx(model: DocumentModel, file_path: Path) -> Path:
    """Save the canonical model to DOCX using semantic headings and list styles."""

    from docx import Document

    doc = Document()
    doc.add_heading(sanitize_inline_text(model.title), level=0)

    for section in model.sections:
        heading_level = max(1, min(9, section.heading_level - 1))
        doc.add_heading(sanitize_inline_text(section.title), level=heading_level)
        for block in section.blocks:
            if block.type == "paragraph":
                doc.add_paragraph(sanitize_inline_text(block.text))
            elif block.type == "bullet_list":
                for item in block.items:
                    doc.add_paragraph(sanitize_inline_text(item), style="List Bullet")
            elif block.type == "numbered_list":
                for item in block.items:
                    doc.add_paragraph(sanitize_inline_text(item), style="List Number")

    parent = file_path.parent.resolve()
    parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(prefix="vaf_", suffix=".docx", dir=str(parent))
    try:
        os.close(fd)
        doc.save(tmp_path)
        if file_path.exists():
            file_path.unlink()
        shutil.move(tmp_path, str(file_path))
        Document(str(file_path))
    finally:
        if os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    return file_path
