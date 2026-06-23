from __future__ import annotations

import base64
from pathlib import Path
from typing import Iterable
from xml.etree.ElementTree import tostring

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from vaf.core.docx_native_model import (
    DocxHeaderFooter,
    DocxImage,
    DocxPageBreak,
    DocxParagraph,
    DocxRun,
    DocxSection,
    DocxSectionProperties,
    DocxTable,
    DocxTableCell,
    DocxTableRow,
    DocxUnsupportedBlock,
    NativeDocxDocument,
    make_warning,
)


_ALIGNMENT_MAP = {
    0: "left",
    1: "center",
    2: "right",
    3: "justify",
}

_SECTION_START_MAP = {
    WD_SECTION_START.NEW_PAGE: "newPage",
    WD_SECTION_START.CONTINUOUS: "continuous",
    WD_SECTION_START.EVEN_PAGE: "evenPage",
    WD_SECTION_START.ODD_PAGE: "oddPage",
    WD_SECTION_START.NEW_COLUMN: "newColumn",
}


def import_docx_to_native_model(path: str | Path) -> NativeDocxDocument:
    target = Path(path).resolve()
    doc = Document(str(target))
    body_children = list(_iter_body_blocks(doc))

    sections: list[DocxSection] = []
    base_sections = list(doc.sections) or [None]

    for section_index, section_ref in enumerate(base_sections):
        sections.append(
            DocxSection(
                id=f"section-{section_index}",
                properties=_section_properties_from_docx(section_ref, section_index),
                header=_header_footer_from_part(getattr(section_ref, "header", None), f"section-{section_index}-header"),
                footer=_header_footer_from_part(getattr(section_ref, "footer", None), f"section-{section_index}-footer"),
                blocks=[],
            )
        )

    if not sections:
        sections = [DocxSection(id="section-0")]

    active_section_index = 0
    for block_index, block in enumerate(body_children):
        current_section = sections[min(active_section_index, len(sections) - 1)]
        if isinstance(block, Paragraph):
            native_block = _paragraph_to_block(doc, block, block_index)
            current_section.blocks.extend(native_block)
            if _paragraph_has_section_break(block) and active_section_index + 1 < len(sections):
                active_section_index += 1
        elif isinstance(block, Table):
            current_section.blocks.append(_table_to_block(block, block_index))
        else:
            if getattr(block, "tag", None) == qn("w:sectPr"):
                # The trailing body <w:sectPr> is the document's final section properties — already captured
                # from doc.sections (_section_properties_from_docx) and re-emitted on export. It is NOT
                # content, so do not turn it into a read-only "Unsupported" placeholder (+ spurious warning).
                continue
            current_section.blocks.append(
                DocxUnsupportedBlock(
                    id=f"unsupported-{block_index}",
                    label="Unsupported body block",
                    xml_tag=str(getattr(block, "tag", "")),
                    xml_payload=tostring(block, encoding="unicode"),
                )
            )

    warnings = _collect_warnings(sections)

    return NativeDocxDocument(
        title=target.stem,
        path=str(target),
        sections=sections,
        warnings=warnings,
    )


def _iter_body_blocks(document: Document) -> Iterable[Paragraph | Table | object]:
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)
        else:
            yield child


def _paragraph_to_block(document: Document, paragraph: Paragraph, block_index: int):
    image_blocks = _extract_image_blocks(document, paragraph, block_index)
    text_runs = [_run_to_native(run, paragraph, block_index, idx) for idx, run in enumerate(paragraph.runs)]
    text_runs = [run for run in text_runs if run.text or run.bold or run.italic or run.underline]

    has_page_break = any(_run_has_page_break(run) for run in paragraph.runs)
    paragraph_block = DocxParagraph(
        id=f"paragraph-{block_index}",
        style_name=str(getattr(paragraph.style, "name", "Normal") or "Normal"),
        alignment=_ALIGNMENT_MAP.get(int(paragraph.alignment)) if paragraph.alignment is not None else "left",
        list_kind=_paragraph_list_kind(paragraph),
        list_level=_paragraph_list_level(paragraph),
        page_break_before=_paragraph_page_break_before(paragraph),
        keep_with_next=_paragraph_keep_with_next(paragraph),
        keep_together=_paragraph_keep_together(paragraph),
        runs=text_runs or [DocxRun(id=f"paragraph-{block_index}-run-0", text=paragraph.text or "")],
    )

    blocks: list = []
    paragraph_text = "".join(run.text for run in paragraph_block.runs).strip()
    if paragraph_text or not image_blocks:
        blocks.append(paragraph_block)
    blocks.extend(image_blocks)
    if has_page_break:
        blocks.append(DocxPageBreak(id=f"page-break-{block_index}"))
    return blocks


def _run_to_native(run, paragraph: Paragraph, block_index: int, run_index: int) -> DocxRun:
    color = ""
    if run.font.color is not None and run.font.color.rgb is not None:
        color = str(run.font.color.rgb)

    highlight = ""
    if run.font.highlight_color is not None:
        highlight = str(run.font.highlight_color)

    font_size_pt = run.font.size.pt if run.font.size is not None else None

    return DocxRun(
        id=f"paragraph-{block_index}-run-{run_index}",
        text=run.text or "",
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        font_name=str(run.font.name or paragraph.style.font.name or ""),
        font_size_pt=float(font_size_pt) if font_size_pt is not None else None,
        color=color,
        highlight=highlight,
    )


def _extract_image_blocks(document: Document, paragraph: Paragraph, block_index: int) -> list[DocxImage]:
    images: list[DocxImage] = []
    for run_index, run in enumerate(paragraph.runs):
        embed_ids = run._element.xpath(".//*[local-name()='blip']/@r:embed")
        if not embed_ids:
            continue
        for image_index, rel_id in enumerate(embed_ids):
            rel = document.part.related_parts.get(rel_id)
            if rel is None:
                continue
            blob = rel.blob
            filename = getattr(rel, "partname", "").split("/")[-1]
            images.append(
                DocxImage(
                    id=f"image-{block_index}-{run_index}-{image_index}",
                    alt_text=paragraph.text.strip() or filename,
                    filename=filename,
                    content_type=str(getattr(rel, "content_type", "")),
                    base64_data=base64.b64encode(blob).decode("ascii"),
                    anchor_kind="inline",
                )
            )
    return images


def _table_to_block(table: Table, block_index: int) -> DocxTable:
    rows: list[DocxTableRow] = []
    for row_index, row in enumerate(table.rows):
        cells: list[DocxTableCell] = []
        for cell_index, cell in enumerate(row.cells):
            paragraphs = [
                DocxParagraph(
                    id=f"table-{block_index}-row-{row_index}-cell-{cell_index}-p-{para_index}",
                    style_name=str(getattr(paragraph.style, "name", "Normal") or "Normal"),
                    alignment=_ALIGNMENT_MAP.get(int(paragraph.alignment)) if paragraph.alignment is not None else "left",
                    list_kind=_paragraph_list_kind(paragraph),
                    list_level=_paragraph_list_level(paragraph),
                    runs=[
                        _run_to_native(run, paragraph, block_index * 1000 + row_index * 100 + cell_index * 10 + para_index, run_index)
                        for run_index, run in enumerate(paragraph.runs)
                    ]
                    or [DocxRun(id=f"table-{block_index}-row-{row_index}-cell-{cell_index}-p-{para_index}-run-0", text=paragraph.text or "")],
                )
                for para_index, paragraph in enumerate(cell.paragraphs)
            ]
            cells.append(
                DocxTableCell(
                    id=f"table-{block_index}-row-{row_index}-cell-{cell_index}",
                    paragraphs=paragraphs or [DocxParagraph(id=f"table-{block_index}-row-{row_index}-cell-{cell_index}-p-0")],
                )
            )
        rows.append(DocxTableRow(id=f"table-{block_index}-row-{row_index}", cells=cells))
    return DocxTable(
        id=f"table-{block_index}",
        rows=rows,
        style_name=str(getattr(table.style, "name", "Table Grid") or "Table Grid"),
    )


def _header_footer_from_part(part, prefix: str) -> DocxHeaderFooter:
    if part is None:
        return DocxHeaderFooter()
    paragraphs = []
    for idx, paragraph in enumerate(getattr(part, "paragraphs", [])):
        text = paragraph.text.strip()
        if not text and not paragraph.runs:
            continue
        paragraphs.append(
            DocxParagraph(
                id=f"{prefix}-paragraph-{idx}",
                style_name=str(getattr(paragraph.style, "name", "Normal") or "Normal"),
                alignment=_ALIGNMENT_MAP.get(int(paragraph.alignment)) if paragraph.alignment is not None else "left",
                runs=[_run_to_native(run, paragraph, idx, run_index) for run_index, run in enumerate(paragraph.runs)]
                or [DocxRun(id=f"{prefix}-paragraph-{idx}-run-0", text=paragraph.text or "")],
            )
        )
    return DocxHeaderFooter(paragraphs=paragraphs)


def _section_properties_from_docx(section_ref, section_index: int) -> DocxSectionProperties:
    if section_ref is None:
        return DocxSectionProperties()
    return DocxSectionProperties(
        page_width_twips=_emu_to_twips(getattr(section_ref, "page_width", None)),
        page_height_twips=_emu_to_twips(getattr(section_ref, "page_height", None)),
        margin_top_twips=_emu_to_twips(getattr(section_ref, "top_margin", None)),
        margin_right_twips=_emu_to_twips(getattr(section_ref, "right_margin", None)),
        margin_bottom_twips=_emu_to_twips(getattr(section_ref, "bottom_margin", None)),
        margin_left_twips=_emu_to_twips(getattr(section_ref, "left_margin", None)),
        start_type=_SECTION_START_MAP.get(getattr(section_ref, "start_type", WD_SECTION_START.NEW_PAGE), "newPage"),
    )


def _emu_to_twips(value) -> int | None:
    if value is None:
        return None
    try:
        return int(value.twips)
    except Exception:
        return None


def _paragraph_has_section_break(paragraph: Paragraph) -> bool:
    return paragraph._element.pPr is not None and paragraph._element.pPr.sectPr is not None


def _paragraph_list_kind(paragraph: Paragraph) -> str:
    style_name = str(getattr(paragraph.style, "name", "") or "").lower()
    if "list bullet" in style_name:
        return "bullet"
    if "list number" in style_name:
        return "numbered"
    num_pr = _paragraph_num_pr(paragraph)
    if num_pr is not None:
        if "bullet" in style_name:
            return "bullet"
        return "numbered"
    return "none"


def _paragraph_list_level(paragraph: Paragraph) -> int:
    num_pr = _paragraph_num_pr(paragraph)
    if num_pr is None:
        return 0
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        return 0
    try:
        return int(ilvl.get(qn("w:val"), "0"))
    except (TypeError, ValueError):
        return 0


def _paragraph_num_pr(paragraph: Paragraph):
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return None
    return p_pr.find(qn("w:numPr"))


def _paragraph_page_break_before(paragraph: Paragraph) -> bool:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return False
    return _on_off_property_enabled(p_pr.find(qn("w:pageBreakBefore")))


def _paragraph_keep_with_next(paragraph: Paragraph) -> bool:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return False
    return _on_off_property_enabled(p_pr.find(qn("w:keepNext")))


def _paragraph_keep_together(paragraph: Paragraph) -> bool:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return False
    return _on_off_property_enabled(p_pr.find(qn("w:keepLines")))


def _on_off_property_enabled(element) -> bool:
    if element is None:
        return False
    raw = element.get(qn("w:val"))
    if raw is None:
        return True
    return str(raw).strip().lower() not in {"0", "false", "off", "no"}


def _run_has_page_break(run) -> bool:
    return bool(run._element.xpath(".//*[local-name()='br' and @w:type='page']"))


def _collect_warnings(sections: list[DocxSection]):
    warnings = []
    for section in sections:
        for block in section.blocks:
            if isinstance(block, DocxUnsupportedBlock):
                warnings.append(
                    make_warning(
                        "unsupported_block",
                        f"Unsupported OOXML block '{block.xml_tag or 'unknown'}' was imported as read-only placeholder.",
                    )
                )
            if isinstance(block, DocxImage) and not block.base64_data:
                warnings.append(
                    make_warning(
                        "image_data_missing",
                        f"Image block '{block.id}' could not be embedded and may not roundtrip correctly.",
                    )
                )
    return warnings
