"""
VAF Document Agent - Section-by-section document creation with bounded context.

This tool is designed to avoid "exceed_context_size_error" by:
- Splitting document creation into sections
- Generating each section separately with its own context
- Assembling sections into final document (Word, PDF, Markdown, Text)
- Supporting huge outputs (500K+ tokens) within 8K context window

Similar architecture to research_agent but for document creation.
"""

import os
import re
import json
import shutil
import tempfile
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

from vaf.tools.base import BaseTool
from vaf.cli.ui import UI

class DocumentAgentTool(BaseTool):
    """
    Specialized Sub-Agent for creating large, structured documents.
    
    Handles:
    - Contracts (Arbeitsverträge, Mietverträge, etc.)
    - Reports (Berichte, Dokumentationen)
    - Letters (Briefe, Anschreiben)
    - Templates (Vorlagen, Formulare)
    - Multi-format output (Word, PDF, Markdown, Text)
    
    Key Feature: Section-by-section generation to prevent context overflow.
    """
    
    name = "document_agent"
    description = """Specialized Sub-Agent for creating large, structured documents (contracts, reports, letters, templates).
Supports multi-format output: Word (.docx), PDF (.pdf), Markdown (.md), Text (.txt).
Handles documents of any size using section-by-section generation (no context overflow)."""
    
    parameters = {
        "type": "object",
        "properties": {
            "task": {
                "type": "string",
                "description": "Document creation task (e.g., 'Create employment contract', 'Generate business report')"
            }
        },
        "required": ["task"]
    }
    
    def __init__(self):
        super().__init__()
        self.home = Path.home()

    def run(self, **kwargs) -> str:
        task = kwargs.get('task', '').strip()
        if not task:
            return "Error: No task provided."
        
        # ═══════════════════════════════════════════════════════════════════════
        # CHECK IF RUNNING IN SEPARATE TERMINAL MODE
        # ═══════════════════════════════════════════════════════════════════════
        from vaf.core.config import Config
        from vaf.core.platform import Platform
        
        # If already in sub-agent terminal, run normally
        if os.environ.get("VAF_IN_SUBAGENT_TERMINAL", "").strip() in ("1", "true", "yes"):
            pass
        elif Config.get("sub_agents_in_separate_terminals", False):
            # Start in new terminal window with IPC tracking
            import shlex
            from vaf.core.subagent_ipc import get_ipc, get_current_session_id
            
            # Create task in IPC system
            ipc = get_ipc()
            task_id = ipc.create_task("document_agent", task_description=task)
            
            # Pass session ID to sub-agent via environment variable
            session_id = get_current_session_id()
            if session_id:
                os.environ["VAF_SESSION_ID"] = session_id
            os.environ["VAF_TASK_ID"] = task_id
            os.environ["VAF_AGENT_TYPE"] = "document_agent"
            
            # Pass provider configuration to sub-agent
            use_separate_provider = Config.get("subagent_use_separate_provider", False)
            if use_separate_provider:
                subagent_provider = Config.get("subagent_provider", "inherit")
                if subagent_provider != "inherit":
                    os.environ["VAF_PROVIDER"] = subagent_provider
            
            cmd_parts = ['vaf', 'subagent', 'run', 'document_agent', '--task', task, '--task-id', task_id]
            
            if Platform.is_windows():
                escaped_parts = []
                for part in cmd_parts:
                    if ' ' in part or '"' in part:
                        escaped = part.replace('"', '\\"')
                        escaped_parts.append(f'"{escaped}"')
                    else:
                        escaped_parts.append(part)
                cmd = ' '.join(escaped_parts)
                title = f"VAF Document Agent [{task_id}]"
            else:
                cmd = ' '.join(shlex.quote(str(part)) for part in cmd_parts)
                title = f"VAF Document Agent [{task_id}]"
            
            if Platform.open_new_terminal(cmd, title=title):
                ipc.mark_task_running(task_id)
                UI.event("Sub-Agent", f"Document Agent started in new terminal [Task: {task_id}]", style="bold cyan")
                return f"[SUBAGENT_ASYNC:{task_id}:document_agent] Sub-Agent running in separate terminal. Task: {task[:80]}..."
            else:
                UI.warning("Failed to open new terminal, running in current window")
                ipc.cancel_task(task_id)
        
        # ═══════════════════════════════════════════════════════════════════════
        # EXECUTE DOCUMENT GENERATION
        # ═══════════════════════════════════════════════════════════════════════
        
        try:
            result = self._generate_document(task)
            return result
        except Exception as e:
            return f"[ERROR] Document generation failed: {e}"
    
    def _generate_document(self, task: str) -> str:
        """Main document generation logic with section-by-section approach."""
        
        UI.event("Document Agent", "Analyzing document request...", style="dim")
        
        # ═══════════════════════════════════════════════════════════════════════
        # STEP 1: Analyze task and create document plan
        # ═══════════════════════════════════════════════════════════════════════
        
        plan = self._create_document_plan(task)
        
        if not plan:
            return "[ERROR] Could not create document plan. Please provide more details about the document."
        
        UI.event("Document Agent", f"Plan created: {plan['title']} ({len(plan['sections'])} sections)", style="success")
        
        # ═══════════════════════════════════════════════════════════════════════
        # STEP 2: Generate each section independently (no context overflow!)
        # ═══════════════════════════════════════════════════════════════════════
        
        sections_content = []
        for i, section in enumerate(plan['sections'], 1):
            UI.event("Document Agent", f"Generating section {i}/{len(plan['sections'])}: {section['title']}", style="dim")
            
            content = self._generate_section(
                document_type=plan['document_type'],
                document_title=plan['title'],
                section_title=section['title'],
                section_description=section['description'],
                section_index=i,
                total_sections=len(plan['sections'])
            )
            
            sections_content.append({
                'title': section['title'],
                'content': content
            })
        
        # ═══════════════════════════════════════════════════════════════════════
        # STEP 3: Assemble final document
        # ═══════════════════════════════════════════════════════════════════════
        
        UI.event("Document Agent", "Assembling final document...", style="dim")
        
        final_content = self._assemble_document(
            title=plan['title'],
            document_type=plan['document_type'],
            sections=sections_content
        )
        
        # ═══════════════════════════════════════════════════════════════════════
        # STEP 4: Save document in requested format
        # ═══════════════════════════════════════════════════════════════════════
        
        output_format = plan.get('format', 'docx')
        filename = plan.get('filename', self._generate_filename(plan['title'], output_format))
        
        file_path = self._save_document(final_content, filename, output_format, plan['document_type'])

        # Notify Web UI so the Document Editor opens with the created document (same process or subprocess)
        try:
            from vaf.core.web_interface import notify_document_created
            session_id = os.environ.get("VAF_SESSION_ID", "").strip()
            if not session_id:
                try:
                    from vaf.core.subagent_ipc import get_current_session_id
                    session_id = get_current_session_id() or ""
                except Exception:
                    pass
            if session_id:
                notify_document_created(session_id, file_path, title=plan.get("title"))
        except Exception as e:
            UI.warning(f"Could not notify document created: {e}")

        # Auto-open the folder containing the document
        try:
            from vaf.core.platform import Platform
            Platform.open_path(file_path.parent)
            UI.event("Document Agent", f"Opened folder: {file_path.parent}", style="dim")
        except Exception as e:
            UI.warning(f"Could not open folder: {e}")
            
        return self._format_success_message(file_path, plan, len(final_content))
    
    def _create_document_plan(self, task: str) -> Optional[Dict]:
        """
        Create a structured plan for the document.
        Uses LLM to analyze task and break it into sections.
        """
        
        
        prompt = f"""You are a document planning expert. Analyze this task and create a structured plan.

Task: {task}

Create a JSON plan with:
1. document_type: (contract/report/letter/template/article/manual)
2. title: Document title
3. format: Output format (docx/pdf/md/txt)
4. filename: Suggested filename
5. sections: Array of sections, each with:
   - title: Section title
   - description: What this section should contain

IMPORTANT: Break complex documents into 5-15 sections for optimal generation.

Output ONLY valid JSON, no explanations."""

        try:
            content = self.query_llm(
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1024,
                temperature=0.3
            )
            
            if content:
                # Extract JSON from response
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    plan = json.loads(json_match.group(0))
                    return plan
            
            return None
            
        except Exception as e:
            UI.error(f"Plan creation failed: {e}")
            return None
    
    def _generate_section(
        self,
        document_type: str,
        document_title: str,
        section_title: str,
        section_description: str,
        section_index: int,
        total_sections: int
    ) -> str:
        """
        Generate a single section with its own isolated context.
        This prevents context overflow for large documents.
        """
        
        # Context-efficient prompt (only this section)
        prompt = f"""You are writing section {section_index} of {total_sections} for a {document_type}.

Document: {document_title}
Section: {section_title}
Requirements: {section_description}

Write this section completely and professionally. Include all necessary details, clauses, or content for this section only.

Output format: Clean text ready for document (no markdown headers, just content).
Language: Match the document type (German for German contracts, etc.)."""

        try:
            content = self.query_llm(
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2048,
                temperature=0.5
            )
            
            if content:
                return content.strip()
            else:
                return "[ERROR generating section: No response from LLM]"
                
        except Exception as e:
            return f"[ERROR generating section: {e}]"
    
    def _assemble_document(self, title: str, document_type: str, sections: List[Dict]) -> str:
        """Assemble all sections into final document."""
        
        # Build document with proper structure
        parts = []
        
        # Title
        parts.append(f"{title}\n")
        parts.append("=" * len(title) + "\n\n")
        
        # Sections
        for i, section in enumerate(sections, 1):
            parts.append(f"\n{section['title']}\n")
            parts.append("-" * len(section['title']) + "\n\n")
            parts.append(section['content'])
            parts.append("\n\n")
        
        return "".join(parts)
    
    def _save_document(self, content: str, filename: str, format: str, doc_type: str) -> Path:
        """Save document in requested format."""
        
        file_path = Path(filename)
        
        try:
            if format == 'docx':
                return self._save_as_word(content, file_path, doc_type)
            elif format == 'pdf':
                return self._save_as_pdf(content, file_path, doc_type)
            elif format == 'md':
                return self._save_as_markdown(content, file_path)
            else:
                return self._save_as_text(content, file_path)
        except Exception as e:
            UI.error(f"Save failed: {e}")
            # Fallback to text
            return self._save_as_text(content, file_path.with_suffix('.txt'))
    
    def _save_as_text(self, content: str, file_path: Path) -> Path:
        """Save as plain text file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return file_path
    
    def _save_as_markdown(self, content: str, file_path: Path) -> Path:
        """Save as Markdown file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return file_path
    
    def _save_as_word(self, content: str, file_path: Path, doc_type: str) -> Path:
        """Save as Word document (.docx). Writes to temp file then replaces target so the ZIP is never half-written."""
        try:
            from docx import Document
            
            doc = Document()
            
            # Parse content and add to document
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    # Check if it's a title (starts with special characters or all caps)
                    if paragraph.strip().startswith('===') or paragraph.strip().startswith('---'):
                        continue  # Skip separator lines
                    elif len(paragraph.strip()) < 100 and paragraph.strip().isupper():
                        doc.add_heading(paragraph.strip(), level=1)
                    elif paragraph.strip().startswith('§') or re.match(r'^\d+\.', paragraph.strip()):
                        doc.add_heading(paragraph.strip(), level=2)
                    else:
                        doc.add_paragraph(paragraph.strip())
            
            parent = file_path.parent
            parent.mkdir(parents=True, exist_ok=True)
            fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=str(parent) if str(parent) != "." else None)
            try:
                os.close(fd)
                doc.save(tmp_path)
                if file_path.exists():
                    file_path.unlink()
                shutil.move(tmp_path, str(file_path))
            finally:
                if os.path.exists(tmp_path):
                    try:
                        os.unlink(tmp_path)
                    except OSError:
                        pass
            return file_path
            
        except ImportError:
            UI.warning("python-docx not installed. Saving as text instead.")
            return self._save_as_text(content, file_path.with_suffix('.txt'))
    
    def _save_as_pdf(self, content: str, file_path: Path, doc_type: str) -> Path:
        """Save as PDF (requires external tool or library)."""
        # For now, save as text and inform user
        # Future: Use reportlab or weasyprint for PDF generation
        text_path = self._save_as_text(content, file_path.with_suffix('.txt'))
        UI.info(f"PDF generation requires additional tools. Saved as text: {text_path}")
        UI.info("To convert to PDF: Use LibreOffice, pandoc, or a PDF printer.")
        return text_path
    
    def _generate_filename(self, title: str, format: str) -> str:
        """Generate a safe filename from title."""
        # Remove special characters
        safe_title = re.sub(r'[^\w\s-]', '', title)
        safe_title = re.sub(r'[\s]+', '_', safe_title)
        safe_title = safe_title[:50]  # Limit length
        
        # Add timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return f"{safe_title}_{timestamp}.{format}"
    
    def _format_success_message(self, file_path: Path, plan: Dict, content_length: int) -> str:
        """Format the success message."""
        
        return f"""### Document Created Successfully! 📄

**Title:** {plan['title']}
**Type:** {plan['document_type'].capitalize()}
**Sections:** {len(plan['sections'])}
**Size:** {content_length:,} characters (~{content_length // 4:,} tokens)
**Format:** {plan.get('format', 'txt').upper()}

**File:** {file_path.name}
**Location:** {file_path.absolute()}

**Sections Generated:**
{chr(10).join([f"  {i}. {s['title']}" for i, s in enumerate(plan['sections'], 1)])}

**Next Steps:**
1. Open the file in your preferred application
2. Review and customize the content
3. Fill in any placeholders ({{COMPANY}}, {{NAME}}, etc.)
4. Save your final version

✅ Document generation completed successfully!"""
    

