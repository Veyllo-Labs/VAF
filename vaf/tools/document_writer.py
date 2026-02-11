"""
Simple Document Writer Tool for Main Agent.

For quick, simple documents (short contracts, letters, messages, templates).
For complex/large documents, use document_agent instead.
"""

import os
import shutil
import tempfile
from pathlib import Path

from vaf.tools.base import BaseTool

class DocumentWriterTool(BaseTool):
    """
    Quick document creation tool for simple documents.
    
    Use for:
    - Short contracts/agreements
    - Letters
    - Messages (WhatsApp, Email templates)
    - Simple forms/templates
    
    For complex or large documents (>5000 chars), use document_agent instead.
    """
    
    name = "document_writer"
    description = """Creates simple structured documents (contracts, letters, messages, templates).
Supports: Text (.txt), Markdown (.md), Word (.docx).
For large/complex documents, use document_agent instead."""
    
    parameters = {
        "type": "object",
        "properties": {
            "document_type": {
                "type": "string",
                "description": "Type of document (e.g., 'contract', 'letter', 'message', 'template')"
            },
            "content": {
                "type": "string",
                "description": "Complete document content (structured text, can use {{PLACEHOLDERS}})"
            },
            "filename": {
                "type": "string",
                "description": "Filename with extension (.txt, .md, .docx)"
            },
            "format": {
                "type": "string",
                "enum": ["text", "markdown", "word"],
                "description": "Output format (default: text)"
            }
        },
        "required": ["document_type", "content", "filename"]
    }
    
    def run(self, **kwargs) -> str:
        document_type = kwargs.get('document_type', 'document')
        content = kwargs.get('content', '')
        filename = kwargs.get('filename', 'document.txt')
        format_type = kwargs.get('format', 'text')
        
        if not content:
            return "[ERROR] No content provided for document."
        
        from vaf.core.platform import Platform
        docs_dir = Platform.documents_dir()
        
        # Create a subdirectory for VAF documents
        vaf_docs_dir = docs_dir / "VAF_Documents"
        vaf_docs_dir.mkdir(exist_ok=True)
        
        # Auto-detect format from filename
        file_path = vaf_docs_dir / Path(filename)
        if file_path.suffix == '.docx':
            format_type = 'word'
        elif file_path.suffix == '.md':
            format_type = 'markdown'
        elif not file_path.suffix:
            file_path = file_path.with_suffix('.txt')
        
        try:
            if format_type == 'word':
                result = self._create_word_document(file_path, content, document_type)
            elif format_type == 'markdown':
                result = self._create_markdown_document(file_path, content, document_type)
            else:
                result = self._create_text_document(file_path, content, document_type)
            # Open the saved document in the Web UI Document Editor
            if not result.startswith("[ERROR]"):
                try:
                    session_id = os.environ.get("VAF_SESSION_ID")
                    if not session_id:
                        from vaf.core.subagent_ipc import get_current_session_id
                        session_id = get_current_session_id()
                    if session_id:
                        from vaf.core.web_interface import notify_document_created
                        notify_document_created(
                            session_id,
                            str(file_path.resolve()),
                            title=file_path.name,
                        )
                except Exception:
                    pass
            return result
        except Exception as e:
            return f"[ERROR] Failed to create document: {e}"
    
    def _create_text_document(self, file_path: Path, content: str, doc_type: str) -> str:
        """Create plain text document."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return f"""### {doc_type.capitalize()} created!

**File:** {file_path.name}
**Path:** {file_path.absolute()}
**Format:** Text
**Size:** {len(content):,} characters

✅ Document saved successfully."""
    
    def _create_markdown_document(self, file_path: Path, content: str, doc_type: str) -> str:
        """Create Markdown document."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return f"""### {doc_type.capitalize()} created!

**File:** {file_path.name}
**Path:** {file_path.absolute()}
**Format:** Markdown
**Size:** {len(content):,} characters

✅ Markdown document saved successfully."""
    
    def _create_word_document(self, file_path: Path, content: str, doc_type: str) -> str:
        """Create Word document (.docx). Writes to temp file then replaces target so the ZIP is never half-written."""
        try:
            from docx import Document
            
            doc = Document()
            
            # Add title
            doc.add_heading(doc_type.capitalize(), 0)
            
            # Split content into paragraphs and add to document
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    # Check if it's a heading (starts with § or #)
                    if paragraph.strip().startswith('§') or paragraph.strip().startswith('#'):
                        doc.add_heading(paragraph.strip().lstrip('#§ '), level=2)
                    else:
                        doc.add_paragraph(paragraph.strip())
            
            parent = file_path.parent
            parent.mkdir(parents=True, exist_ok=True)
            fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=str(parent))
            try:
                os.close(fd)
                doc.save(tmp_path)
                if file_path.exists():
                    file_path.unlink()
                shutil.move(tmp_path, str(file_path))
            finally:
                if os.path.exists(tmp_path):
                    try:
                        os.unlink(tmp_path)
                    except OSError:
                        pass
            
            return f"""### {doc_type.capitalize()} created!

**File:** {file_path.name}
**Path:** {file_path.absolute()}
**Format:** Microsoft Word (.docx)
**Size:** {len(content):,} characters

✅ Word document saved successfully.
   Open with: Microsoft Word, LibreOffice, Google Docs"""
            
        except ImportError:
            return f"""[ERROR] python-docx not installed.

To create Word documents, run:
    pip install python-docx

Alternative: Save as text (.txt) or Markdown (.md) instead."""
        except Exception as e:
            return f"[ERROR] Failed to create Word document: {e}"
